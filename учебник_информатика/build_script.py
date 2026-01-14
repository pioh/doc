#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт сборки учебника по информатике в формат DOCX

Использование:
    python build_script.py

Результат:
    Создается файл учебник_информатика.docx в текущей директории
"""

import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("❌ ОШИБКА: Необходимо установить библиотеку python-docx")
    print("Выполните: pip install python-docx")
    exit(1)


class TextbookBuilder:
    """Класс для сборки учебника из markdown файлов в DOCX"""
    
    def __init__(self, chapters_dir="chapters", output_file="учебник_информатика.docx"):
        self.chapters_dir = Path(chapters_dir)
        self.output_file = output_file
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Настройка стилей документа"""
        # Стиль для обычного текста
        styles = self.doc.styles
        style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Настройка параграфа
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_after = Pt(6)
        
    def add_title_page(self):
        """Добавление титульной страницы"""
        # Заголовок
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("УЧЕБНИК ПО ИНФОРМАТИКЕ\n\n")
        run.font.size = Pt(18)
        run.font.bold = True
        
        # Подзаголовок
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Для студентов первого курса\n\n")
        run.font.size = Pt(14)
        
        # Информация
        info = self.doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(f"\n\n\n\n{datetime.now().year}")
        run.font.size = Pt(12)
        
        self.doc.add_page_break()
        
    def add_table_of_contents(self):
        """Добавление оглавления"""
        heading = self.doc.add_heading("ОГЛАВЛЕНИЕ", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            with open("table_of_contents.md", "r", encoding="utf-8") as f:
                content = f.read()
                
            # Парсим оглавление
            lines = content.split('\n')
            for line in lines:
                if line.startswith('## Раздел'):
                    p = self.doc.add_paragraph(line.replace('## ', ''))
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(14)
                elif line.startswith('**Глава'):
                    # Извлекаем номер и название главы
                    match = re.match(r'\*\*Глава ([\d.]+)\*\* (.+)', line)
                    if match:
                        chapter_num = match.group(1)
                        chapter_name = match.group(2)
                        p = self.doc.add_paragraph(f"Глава {chapter_num} {chapter_name}")
                        p.paragraph_format.left_indent = Inches(0.5)
                        
        except FileNotFoundError:
            print("⚠️ Предупреждение: файл table_of_contents.md не найден")
            
        self.doc.add_page_break()
        
    def get_chapter_files(self):
        """Получение списка файлов глав в правильном порядке"""
        if not self.chapters_dir.exists():
            print(f"❌ ОШИБКА: Папка {self.chapters_dir} не найдена!")
            return []
            
        # Получаем все .md файлы и сортируем их
        chapter_files = sorted(self.chapters_dir.glob("*.md"))
        return chapter_files
        
    def parse_markdown_line(self, line, paragraph=None):
        """Парсинг одной строки markdown"""
        if paragraph is None:
            paragraph = self.doc.add_paragraph()
            
        # Обработка жирного текста **text**
        parts = re.split(r'(\*\*[^*]+\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part:
                paragraph.add_run(part)
                
        return paragraph
        
    def add_chapter(self, chapter_file):
        """Добавление главы в документ"""
        print(f"  📄 Обработка: {chapter_file.name}")
        
        try:
            with open(chapter_file, "r", encoding="utf-8") as f:
                content = f.read()
        except Exception as e:
            print(f"  ⚠️ Ошибка чтения файла: {e}")
            return
            
        # Проверка на пустой файл
        if not content.strip() or content.strip() == "# TODO: Эта глава будет написана":
            print(f"  ⏭️ Пропуск: глава еще не написана")
            return
            
        lines = content.split('\n')
        in_code_block = False
        code_lines = []
        
        for line in lines:
            # Обработка блоков кода
            if line.startswith('```'):
                if in_code_block:
                    # Конец блока кода
                    code_text = '\n'.join(code_lines)
                    p = self.doc.add_paragraph(code_text)
                    p.style = 'No Spacing'
                    p.runs[0].font.name = 'Courier New'
                    p.runs[0].font.size = Pt(10)
                    code_lines = []
                in_code_block = not in_code_block
                continue
                
            if in_code_block:
                code_lines.append(line)
                continue
                
            # Заголовки
            if line.startswith('# '):
                heading = self.doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                heading = self.doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                heading = self.doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                heading = self.doc.add_heading(line[5:], level=4)
            # Списки
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                p = self.doc.add_paragraph(text, style='List Bullet')
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = self.doc.add_paragraph(text, style='List Number')
            # Пустые строки
            elif not line.strip():
                self.doc.add_paragraph()
            # Обычный текст
            else:
                self.parse_markdown_line(line)
                
        # Разрыв страницы после главы
        self.doc.add_page_break()
        
    def build(self):
        """Основной метод сборки"""
        print("=" * 60)
        print("🚀 СБОРКА УЧЕБНИКА ПО ИНФОРМАТИКЕ")
        print("=" * 60)
        
        # Добавление титульной страницы
        print("\n📋 Создание титульной страницы...")
        self.add_title_page()
        
        # Добавление оглавления
        print("📋 Создание оглавления...")
        self.add_table_of_contents()
        
        # Получение списка глав
        chapter_files = self.get_chapter_files()
        
        if not chapter_files:
            print("❌ ОШИБКА: Главы не найдены!")
            return False
            
        print(f"\n📚 Найдено глав: {len(chapter_files)}")
        print("\n📝 Обработка глав:")
        
        # Добавление всех глав
        chapters_added = 0
        for chapter_file in chapter_files:
            self.add_chapter(chapter_file)
            chapters_added += 1
            
        # Сохранение документа
        print(f"\n💾 Сохранение документа: {self.output_file}")
        try:
            self.doc.save(self.output_file)
            print(f"✅ УСПЕШНО! Документ сохранен: {self.output_file}")
            print(f"📊 Обработано глав: {chapters_added}")
            return True
        except Exception as e:
            print(f"❌ ОШИБКА при сохранении: {e}")
            return False


def main():
    """Главная функция"""
    # Проверка текущей директории
    if not Path("spec.md").exists():
        print("❌ ОШИБКА: Запустите скрипт из папки учебник_информатика/")
        return
        
    # Создание и запуск сборщика
    builder = TextbookBuilder()
    success = builder.build()
    
    if success:
        print("\n" + "=" * 60)
        print("🎉 СБОРКА ЗАВЕРШЕНА УСПЕШНО!")
        print("=" * 60)
        print(f"\n📖 Ваш учебник готов: {builder.output_file}")
        print("🖨️  Можно отправлять на печать!")
    else:
        print("\n" + "=" * 60)
        print("❌ СБОРКА ЗАВЕРШИЛАСЬ С ОШИБКАМИ")
        print("=" * 60)


if __name__ == "__main__":
    main()
